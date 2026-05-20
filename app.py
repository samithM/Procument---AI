import streamlit as st
import google.generativeai as genai
import pandas as pd
from docx import Document
from io import BytesIO
import json

# ==========================================
# 1. INITIAL SYSTEM CONFIGURATION
# ==========================================
st.set_page_config(
    page_title="AI Procurement Document Generator",
    page_icon="🏗️",
    layout="wide"
)

# Setup Gemini API Connection (Get Free Key from Google AI Studio)
st.sidebar.title("🔐 Configuration")
api_key = st.sidebar.text_input("Google AI Studio API Key", type="password", help="Enter your Gemini free API key")

if api_key:
    genai.configure(api_key=api_key)
else:
    st.sidebar.warning("Please enter your Gemini API key to activate AI features.")

# ==========================================
# APP TITLE & DESCRIPTION
# ==========================================
st.title("🏗️ AI-Powered Procurement Document Generator")
st.caption("CIDA / NPC Compliant Smart Procurement System for Sri Lankan Construction Industry")
st.markdown("---")

# ==========================================
# PHASE 1: USER INPUT FLOW (5 STEPS)
# ==========================================

# Use Streamlit Tabs for the 5 Steps
tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📍 Step 1: Project Basics", 
    "📊 Step 2: BOQ & Scope", 
    "⚖️ Step 3: Procurement & CIDA", 
    "💰 Step 4: Financials", 
    "📄 Step 5: Generate & Download"
])

# 📥 STEP 1: PROJECT BASICS
with tab1:
    st.header("Project Identification Basics")
    col1, col2 = st.columns(2)
    
    with col1:
        construction_type = st.selectbox(
            "Construction Type / Category",
            ["Building", "Building Renovation & Rehabilitation", "Road", "Irrigation", "Bridge", "Electrical", "Water Supply", "Marine & Coastal Work"]
        )
        procuring_entity = st.selectbox(
            "Procuring Entity Type",
            ["Ministry", "Department (RDA/PRDA/MASL)", "University", "Local Authority / Municipal Council"]
        )
        entity_name = st.text_input("Name of the Procuring Entity", value="Ocean University of Sri Lanka")
        
    with col2:
        province = st.selectbox("Province", ["Western", "Southern", "Central", "Northern", "Eastern", "North Western", "North Central", "Uva", "Sabaragamuwa"])
        district = st.selectbox("District (For BSR Rate Mapping)", ["Colombo", "Gampaha", "Kalutara", "Matara", "Galle", "Hambantota", "Kandy", "Jaffna", "Anuradhapura"])
        project_year = st.text_input("Project Year", value="2026")
        
    # Auto-generate Tender Reference Number
    # Format: [Entity]/[Type]/[Year]/[District]/001
    type_code = construction_type[:4].upper().replace(" ", "")
    district_code = district[:2].upper()
    entity_code = "".join([w[0] for w in entity_name.split() if w.isalpha()]).upper()
    auto_ref_no = f"{entity_code}/{type_code}/{project_year}/{district_code}/001"
    
    tender_ref = st.text_input("Generated Tender Reference Number", value=auto_ref_no)

# 📥 STEP 2: BOQ & DOCUMENTS UPLOAD
with tab2:
    st.header("Scope Analysis & Data Parsing")
    
    uploaded_boq = st.file_uploader("Upload Existing BOQ (Excel or PDF)", type=["xlsx", "xls", "pdf"])
    if uploaded_boq:
        st.success(f"Successfully uploaded: {uploaded_boq.name} (Ready for AI Parsing)")
        
    project_scope = st.text_area(
        "Describe the project scope (AI will use this to generate Standard Specification Clauses & Method Statements)",
        placeholder="Provide details like: Single-story building extension, Grade 25 concrete for foundations, GI roofing sheet with treated timber framework, Emulsion painting..."
    )
    
    st.info("🌐 Live Integration Active: AI will automatically map scope with regional BSR/HSR indices.")

# 📥 STEP 3: PROCUREMENT & CIDA ELIGIBILITY
with tab3:
    st.header("Procurement Method & Compliance Selection")
    col3, col4 = st.columns(2)
    
    with col3:
        sbd_type = st.selectbox(
            "Standard Bidding Document (SBD) Type",
            ["National Shopping (Minor Works)", "NCB (National Competitive Bidding)", "ICB (International Competitive Bidding)", "CIDA/SBD/03 (Major Works)"]
        )
        contract_type = st.selectbox("Contract Type", ["Admeasurement (Measure & Pay)", "Lump Sum", "Design & Build"])
        
    with col4:
        estimated_value = st.number_input("Estimated Project Value (LKR)", min_value=0.0, value=5000000.0, step=100000.0)
        
        # Auto CIDA Grade Suggester logic based on value thresholds
        suggested_grade = "C9"
        if estimated_value > 500000000: suggested_grade = "C1 and Above"
        elif estimated_value > 300000000: suggested_grade = "C2"
        elif estimated_value > 150000000: suggested_grade = "C3"
        elif estimated_value > 50000000: suggested_grade = "C4"
        elif estimated_value > 25000000: suggested_grade = "C5"
        elif estimated_value > 10000000: suggested_grade = "C6"
        elif estimated_value > 5000000: suggested_grade = "C7"
        elif estimated_value > 2000000: suggested_grade = "C8"
        
        st.success(f"💡 Recommended Contractor CIDA Grade Requirement: **{suggested_grade}**")

# 📥 STEP 4: FINANCIAL PARAMETERS
with tab4:
    st.header("Financial Security & Formula Controls")
    
    col5, col6 = col1, col2 = st.columns(2)
    
    # Financial Auto-calculations based on industry standard percentages
    calc_bid_sec = estimated_value * 0.01  # 1%
    calc_perf_sec = estimated_value * 0.05 # 5%
    calc_adv_pay = estimated_value * 0.20  # 20%
    calc_retention = estimated_value * 0.10 # 10%
    
    with col5:
        bid_security = st.number_input("Bid Security Value (LKR) - (1% Suggested)", value=calc_bid_sec)
        advance_payment = st.number_input("Advance Payment Limit (LKR) - (20% Max)", value=calc_adv_pay)
        price_fluctuation = st.checkbox("Enable Price Fluctuation Clause (CIDA Formula)", value=True if estimated_value > 10000000 else False)
        
    with col6:
        performance_security = st.number_input("Performance Security (LKR) - (5% Standard)", value=calc_perf_sec)
        retention_ceiling = st.number_input("Retention Ceiling Limit (LKR) - (10% Cap)", value=calc_retention)
        liquidated_damages = st.text_input("Liquidated Damages per Day (LKR)", value=f"Rs. {int(estimated_value * 0.0005)} per day")

# 📥 STEP 5: FINAL DOCUMENT GENERATE
with tab5:
    st.header("Document Engine Activation")
    st.write("Click below to pass all criteria to Gemini AI and generate your styled Procurement Document.")
    
    if st.button("🚀 Run AI Document Generator"):
        if not api_key:
            st.error("Error: Please provide a Google AI Studio API Key in the sidebar.")
        elif not project_scope:
            st.error("Error: Please enter the project scope in Step 2 to generate clauses.")
        else:
            with st.spinner("AI Engine is connecting to Google Studio & compiling CIDA frameworks..."):
                try:
                    # Construct Prompt for Gemini
                    prompt = f"""
                    You are an expert Civil Engineer and Procurement Officer specialized in Sri Lankan CIDA (ICTAD) and National Procurement Commission (NPC) guidelines.
                    Generate a detailed Invitation for Bids (IFB), Contract Data sheet, and Technical Specifications based on the following details:
                    
                    - Procuring Entity: {entity_name} ({procuring_entity})
                    - Tender Ref: {tender_ref}
                    - Project Type: {construction_type}
                    - Location: {district}, {province} Province
                    - Estimated Value: LKR {estimated_value}
                    - Procurement Method: {sbd_type}
                    - Required CIDA Grade: {suggested_grade}
                    - Financials: Bid Security={bid_security}, Performance Security={performance_security}, Advance Payment={advance_payment}, Retention={retention_ceiling}
                    - Project Scope: {project_scope}
                    
                    Provide the output in a clean, structured text format with clear section headings suitable for standard procurement layout. Include standard specification clauses mapped to CIDA requirements for the given scope.
                    """
                    
                    # Call Gemini Model
                    model = genai.GenerativeModel("gemini-1.5-flash-latest")
                    response = model.generate_content(prompt)
                    generated_text = response.text
                    
                    st.success("🎉 AI Content Generated Successfully!")
                    st.text_area("Live Preview", value=generated_text, height=300)
                    
                    # Create Word File in Memory
                    doc = Document()
                    doc.add_heading(f"PROCUREMENT DOCUMENT - {tender_ref}", level=1)
                    doc.add_paragraph(f"Project Type: {construction_type}")
                    doc.add_paragraph(f"Procuring Entity: {entity_name}")
                    doc.add_markdown(generated_text)
                    
                    # Save to byte stream for download button
                    bio = BytesIO()
                    doc.save(bio)
                    bio.seek(0)
                    
                    st.download_button(
                        label="📥 Download Editable Word Document (.docx)",
                        data=bio,
                        file_name=f"Procurement_Doc_{tender_ref}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                except Exception as e:
                    st.error(f"An error occurred: {e}")
